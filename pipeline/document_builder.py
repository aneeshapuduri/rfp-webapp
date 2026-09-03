"""
Phase 4: assembles the full proposal into a .docx following template_mapper.OUR_SECTIONS'
canonical order (cover, table of contents, then every section from Executive Summary and
Bidder Overview through Closing Statement) — a mix of content Phases 1-4 generated
(narrative, timeline, staffing, pricing, compliance matrix) and static content pulled
straight from config/company_profile.json (company overview, value proposition, core
capabilities, differentiators, partnerships, industry recognition, key personnel, SLA, etc.).

Section content is produced by a registry of small renderer functions (SECTION_RENDERERS),
one per canonical section name, each called as renderer(sink, content, company). A "sink" is
one of two interchangeable output targets: _AppendSink writes at the current end of the
document (used by the default no-template builder, and for template sections that couldn't be
matched to a heading), while _InsertAfterSink writes immediately after a moving cursor anchored
at a heading found in a client-supplied template. Renderers never need to know which sink
they're writing to — this is what let the old 3x-duplicated per-section if/elif chains (one
copy each in the default builder, the matched-template path, and the unmatched-append path)
collapse into a single implementation per section.
"""
from __future__ import annotations

import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from template_mapper import OUR_SECTIONS

NAVY = RGBColor(0x1F, 0x2D, 0x50)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
GRAY = RGBColor(0x59, 0x59, 0x59)
GREEN = "D9EAD3"
ORANGE = "FCE5CD"


def _shade(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tc_pr.append(shd)


def _style_base(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for level, size in [(1, 18), (2, 14)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = NAVY
        style.font.bold = True
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(8)


def _add_footer(doc: Document, project_title: str):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = f"{project_title} — Confidential — Proposal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY


def _add_cover_page(doc: Document, project_title: str, agency: str, company: dict):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPOSAL")
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    run.font.bold = True

    proj = doc.add_paragraph()
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj.add_run(project_title)
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Submitted in response to a Request for Proposal from {agency}")
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY
    run.italic = True

    for _ in range(6):
        doc.add_paragraph()

    for text, size, bold, italic in [
        (company.get("company_name", ""), 15, True, False),
        (company.get("tagline", ""), 11, False, True),
        (company.get("hq_location", ""), 10, False, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.italic = italic
        run.font.color.rgb = NAVY if bold else GRAY

    for _ in range(3):
        doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.date.today().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_toc(doc: Document):
    """Inserts a real Word Table-of-Contents field (TOC \\o "1-2" \\h \\z \\u) on its own page
    right after the cover page. python-docx cannot compute page numbers itself, so the field
    renders with placeholder text until opened in Word — Word normally prompts to update fields
    on open, and it can always be refreshed manually (right-click inside it -> Update Field, or
    select it and press F9)."""
    doc.add_heading("Table of Contents", level=1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r_element = run._r

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = 'Right-click here and choose "Update Field" (or select it and press F9) to generate the table of contents.'

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(placeholder)
    r_element.append(fld_char_end)

    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_body(doc: Document, text: str):
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float], status_col: int | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(w) for w in widths_in]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].width = widths[i]
        _shade(hdr[i], "1F2D50")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].width = widths[i]
        if status_col is not None:
            status_val = str(row_data[status_col])
            _shade(row[status_col], GREEN if "Full" in status_val else ORANGE)
    return table


def _find_heading_element(doc: Document, heading_text: str):
    """Locates a paragraph in an opened (client-supplied) template whose style is a Heading
    style and whose text matches exactly — this is always a heading text returned by
    template_mapper.extract_template_headings() on this same document, so an exact match is
    expected. Returns the underlying XML element (not the python-docx Paragraph wrapper) since
    that's what the insertion helpers below operate on."""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip() == heading_text:
            return p._p
    return None


def _insert_paragraph_after(doc: Document, anchor_element, text: str = "", style: str | None = None):
    """python-docx only ever appends new paragraphs/tables at the very end of the document body.
    To place generated content immediately under an existing heading somewhere in the middle of
    a client's uploaded template, this creates the paragraph normally (so it lands at the end),
    then relocates its underlying XML element to sit right after anchor_element. Returns the new
    element so a caller inserting multiple lines in a row can chain them in order."""
    new_p = doc.add_paragraph(text, style=style)
    anchor_element.addnext(new_p._p)
    return new_p._p


def _insert_body_after(doc: Document, anchor_element, text: str):
    cursor = anchor_element
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("- "):
            cursor = _insert_paragraph_after(doc, cursor, line[2:], style="List Bullet")
        else:
            cursor = _insert_paragraph_after(doc, cursor, line)
    return cursor


def _insert_table_after(doc: Document, anchor_element, headers, rows, widths_in, status_col=None):
    table = _add_table(doc, headers, rows, widths_in, status_col)  # appended at document end
    anchor_element.addnext(table._tbl)
    return table._tbl


class _AppendSink:
    """Writes content at the current end of the document. Used by the default (no-template)
    builder, and for sections that couldn't be confidently matched to a heading in a client's
    uploaded template (the 'Needs Manual Placement' block)."""

    def __init__(self, doc: Document):
        self.doc = doc

    def body(self, text: str):
        _add_body(self.doc, text)

    def bullets(self, items: list[str], empty_text: str | None = None):
        if items:
            for item in items:
                self.doc.add_paragraph(item, style="List Bullet")
        elif empty_text:
            self.doc.add_paragraph(empty_text)

    def table(self, headers, rows, widths_in, status_col=None):
        return _add_table(self.doc, headers, rows, widths_in, status_col=status_col)

    def paragraph(self, text: str = "", bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None, size: int | None = None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
        return p


class _InsertAfterSink:
    """Writes content immediately after a moving cursor anchored at a heading found in a
    client-supplied template (see _insert_paragraph_after / _insert_table_after). Mirrors
    _AppendSink's interface so a SECTION_RENDERERS function never needs to know which case
    it's writing into."""

    def __init__(self, doc: Document, anchor_element):
        self.doc = doc
        self.cursor = anchor_element

    def body(self, text: str):
        self.cursor = _insert_body_after(self.doc, self.cursor, text)

    def bullets(self, items: list[str], empty_text: str | None = None):
        if items:
            for item in items:
                self.cursor = _insert_paragraph_after(self.doc, self.cursor, item, style="List Bullet")
        elif empty_text:
            self.cursor = _insert_paragraph_after(self.doc, self.cursor, empty_text)

    def table(self, headers, rows, widths_in, status_col=None):
        self.cursor = _insert_table_after(self.doc, self.cursor, headers, rows, widths_in, status_col=status_col)
        return self.cursor

    def paragraph(self, text: str = "", bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None, size: int | None = None):
        new_p = self.doc.add_paragraph()
        run = new_p.add_run(text)
        run.font.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
        self.cursor.addnext(new_p._p)
        self.cursor = new_p._p
        return new_p


def _render_executive_summary(sink, content, company):
    sink.body(content["executive_summary"])


def _render_understanding(sink, content, company):
    sink.body(content["understanding"])


def _render_company_overview(sink, content, company):
    if company.get("company_name"):
        sink.body(
            f"{company['company_name']} was founded in {company.get('founded_year', 'N/A')} and is "
            f"headquartered in {company.get('hq_location', 'N/A')}, with {company.get('employee_count', 'N/A')} "
            f"employees. {company.get('tagline', '')}"
        )
    certs = company.get("certifications", [])
    if certs:
        sink.bullets(certs)


def _render_value_proposition(sink, content, company):
    sink.body(company.get("value_proposition", ""))


def _render_core_capabilities(sink, content, company):
    sink.bullets(company.get("core_capabilities", []))


def _render_technologies_and_skillsets(sink, content, company):
    sink.bullets(company.get("technologies_and_skillsets", []))


def _render_key_differentiators(sink, content, company):
    sink.bullets(company.get("differentiators", []))


def _render_partnership_list(sink, content, company):
    partnerships = company.get("partnerships", [])
    if not partnerships:
        sink.body("No formal technology partnerships apply to this engagement.")
        return
    sink.bullets([f"{p['name']} — {p['description']}" for p in partnerships])


def _render_industry_recognition(sink, content, company):
    sink.bullets(
        company.get("industry_recognition", []),
        empty_text="No industry recognition to report at this time.",
    )


def _render_past_performance(sink, content, company):
    sink.body(content["past_performance"])


def _render_staffing_readiness(sink, content, company):
    sink.body(content.get("staffing_readiness", ""))


def _render_scope_of_services(sink, content, company):
    sink.body(content.get("scope_of_services", ""))


def _render_out_of_scope(sink, content, company):
    sink.body(content.get("out_of_scope_of_services", ""))


def _render_project_objectives(sink, content, company):
    sink.body(content.get("project_objectives", ""))


def _render_project_deliverables(sink, content, company):
    sink.body(content.get("project_deliverables", ""))


def _render_solution_overview(sink, content, company):
    sink.body(content["technology_approach"])


def _render_project_plan_milestones(sink, content, company):
    timeline_rows = [[p["phase"], p["duration"], p["description"]] for p in content["timeline"]]
    sink.table(["Phase", "Duration", "Description"], timeline_rows, [1.6, 1.3, 4.1])


def _render_change_management(sink, content, company):
    sink.body(company.get("change_management_approach", ""))


def _render_performance_optimization(sink, content, company):
    sink.body(company.get("performance_optimization_approach", ""))


def _render_key_personnel(sink, content, company):
    for person in company.get("key_personnel", []):
        sink.paragraph(f"{person['name']} — {person['role']}", bold=True, color=NAVY)
        sink.paragraph(person.get("credentials", ""), italic=True, color=GRAY)
        sink.paragraph(person.get("bio", ""))


def _render_team_staffing_plan(sink, content, company):
    staffing_rows = [
        [l["role"], l["headcount"], l["hours_per_person"], l["total_hours"]]
        for l in content["staffing"]
    ]
    sink.table(["Role", "Headcount", "Hours/Person", "Total Hours"], staffing_rows, [2.5, 1.2, 1.4, 1.4])


def _render_maintenance_and_support(sink, content, company):
    sink.body(company.get("maintenance_and_support", ""))


def _render_operating_support_model(sink, content, company):
    sink.body(company.get("operating_support_model", ""))


def _render_technical_assumptions(sink, content, company):
    sink.body(content.get("technical_assumptions", ""))


def _render_general_assumptions(sink, content, company):
    sink.bullets(
        content["assumptions"],
        empty_text="No assumptions were required — all requirements were fully specified in the RFP.",
    )


def _render_project_dependencies(sink, content, company):
    sink.body(content.get("project_dependencies", ""))


def _render_pricing_summary(sink, content, company):
    pricing = content["pricing"]
    pricing_rows = [
        [l["role"], l["total_hours"], f"${l['hourly_rate']:,.2f}", f"${l['subtotal']:,.2f}"]
        for l in pricing["lines"]
    ]
    sink.table(["Role", "Total Hours", "Rate", "Subtotal"], pricing_rows, [2.3, 1.3, 1.3, 1.6])
    sink.paragraph(f"Labor Subtotal: ${pricing['labor_subtotal']:,.2f}")
    sink.paragraph(f"Contingency ({pricing['contingency_pct']}%): ${pricing['contingency_amount']:,.2f}")
    sink.paragraph(f"TOTAL: ${pricing['total']:,.2f}", bold=True, size=13, color=NAVY)


def _render_sla(sink, content, company):
    sla_rows = [[s["severity"], s["response_time"], s["resolution_time"]] for s in company.get("sla_default", [])]
    if sla_rows:
        sink.table(["Severity", "Response Time", "Resolution Time"], sla_rows, [2.6, 1.9, 1.9])
    else:
        sink.body("Service levels will be finalized in the executed agreement.")


def _render_service_boundaries(sink, content, company):
    sink.body(content.get("service_boundaries", ""))


def _render_compliance_matrix(sink, content, company):
    matrix_rows = [[m["requirement"], m["response"], m["status"]] for m in content["compliance_matrix"]]
    sink.table(["Requirement", "Our Response", "Status"], matrix_rows, [2.3, 3.4, 1.3], status_col=2)


def _render_closing(sink, content, company):
    sink.body(content["closing"])


# One renderer per canonical section in template_mapper.OUR_SECTIONS. A renderer is called as
# renderer(sink, content, company) and writes that section's body content through the sink —
# it never adds the section's own heading (the caller does, since a matched-in-template section
# reuses the heading already present in the client's document) and never needs to know whether
# it's appending at the document's end or inserting after a cursor in the middle of one.
SECTION_RENDERERS = {
    "Executive Summary and Bidder Overview": _render_executive_summary,
    "Understanding of the Problem": _render_understanding,
    "Company Overview": _render_company_overview,
    "Company Value Proposition": _render_value_proposition,
    "Core Capabilities": _render_core_capabilities,
    "Technologies and Skillsets": _render_technologies_and_skillsets,
    "Company's Key Differentiators": _render_key_differentiators,
    "Partnership List": _render_partnership_list,
    "Industry Recognition": _render_industry_recognition,
    "Relevant Past Performance": _render_past_performance,
    "Staffing Model and Project Initiation Readiness": _render_staffing_readiness,
    "Scope of Services": _render_scope_of_services,
    "Out of Scope of Services": _render_out_of_scope,
    "Project Objectives": _render_project_objectives,
    "Project Deliverables": _render_project_deliverables,
    "High-Level Solution Overview": _render_solution_overview,
    "Project Plan and Milestones": _render_project_plan_milestones,
    "Project Change Management": _render_change_management,
    "Performance Optimization Approach": _render_performance_optimization,
    "Key Personnel": _render_key_personnel,
    "Proposed Team & Staffing Plan": _render_team_staffing_plan,
    "Maintenance and Support": _render_maintenance_and_support,
    "Operating Support Model": _render_operating_support_model,
    "Technical Assumptions": _render_technical_assumptions,
    "General Assumptions": _render_general_assumptions,
    "Project Assumptions and Dependencies": _render_project_dependencies,
    "Effort & Pricing Summary": _render_pricing_summary,
    "Service Level Agreement (SLA)": _render_sla,
    "Service Boundaries and Scope Limitations": _render_service_boundaries,
    "Compliance Matrix": _render_compliance_matrix,
    "Closing Statement": _render_closing,
}


def _insert_matched_section(doc: Document, our_section: str, template_heading: str, content: dict, company: dict):
    """Inserts generated content for a canonical section immediately after the matching heading
    found in the client's uploaded template. No section numbering is added here — the client's
    own template heading is used exactly as they wrote it."""
    anchor = _find_heading_element(doc, template_heading)
    renderer = SECTION_RENDERERS.get(our_section)
    if anchor is None or renderer is None:
        return
    renderer(_InsertAfterSink(doc, anchor), content, company)


def _append_section_content(doc: Document, our_section: str, content: dict, company: dict):
    """Same section content as _insert_matched_section, but appended normally at the end of the
    document — used for sections that couldn't be confidently matched to a heading in the
    client's template (see the 'Needs Manual Placement' block in build_final_proposal)."""
    renderer = SECTION_RENDERERS.get(our_section)
    if renderer is None:
        return
    renderer(_AppendSink(doc), content, company)


def _build_final_proposal_with_template(output_path: str, content: dict, template_path: str,
                                         section_mapping: dict | None):
    """Writes generated content into the client's own uploaded .docx template instead of a
    fresh document — deliberately does NOT call _style_base or _add_cover_page or _add_footer:
    the whole point of a custom template is to keep the client's own styling and cover/branding
    intact, so this only ever inserts new content, never restyles or replaces what's already
    there. Any canonical section template_mapper couldn't confidently match to a heading in the
    template is appended at the end under a 'Needs Manual Placement' heading instead of being
    silently dropped."""
    company = content["company"]
    doc = Document(template_path)

    mapping = section_mapping or {"matched": {}, "unmatched": list(OUR_SECTIONS)}
    for our_section, template_heading in mapping.get("matched", {}).items():
        _insert_matched_section(doc, our_section, template_heading, content, company)

    unmatched = mapping.get("unmatched", [])
    if unmatched:
        doc.add_heading("Needs Manual Placement", level=1)
        note = doc.add_paragraph()
        note.add_run(
            "The following sections could not be confidently matched to a heading in your "
            "uploaded template and have been placed here instead — please move each into the "
            "appropriate place in the document."
        ).italic = True
        for our_section in unmatched:
            doc.add_heading(our_section, level=2)
            _append_section_content(doc, our_section, content, company)

    doc.save(output_path)
    return output_path


def build_final_proposal(output_path: str, content: dict, template_path: str | None = None,
                          section_mapping: dict | None = None):
    """content: the dict returned by phase4_pipeline.run_phase4() (optionally edited by the user
    in the preview stage). When template_path is given (a client-uploaded .docx), section_mapping
    should be the dict returned by template_mapper.map_sections_to_template() for that same
    template's headings — content is then written into the client's own document instead of a
    fresh one (see _build_final_proposal_with_template). template_path is None reproduces the
    original default-template behavior exactly as before this parameter existed."""
    if template_path is not None:
        return _build_final_proposal_with_template(output_path, content, template_path, section_mapping)

    company = content["company"]
    doc = Document()
    _style_base(doc)
    _add_cover_page(doc, content["project_title"], content["agency"], company)
    _add_toc(doc)
    _add_footer(doc, content["project_title"])

    sink = _AppendSink(doc)
    for i, section in enumerate(OUR_SECTIONS, start=1):
        doc.add_heading(f"{i}. {section}", level=1)
        renderer = SECTION_RENDERERS.get(section)
        if renderer is not None:
            renderer(sink, content, company)

    contact = company.get("contact", {})
    if contact:
        doc.add_paragraph()
        cp = doc.add_paragraph()
        run = cp.add_run(
            f"{contact.get('name', '')}, {contact.get('title', '')}  |  "
            f"{contact.get('email', '')}  |  {contact.get('phone', '')}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    doc.save(output_path)
    return output_path
