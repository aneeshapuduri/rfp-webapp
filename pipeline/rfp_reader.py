"""
Reads an RFP document (.txt, .docx, or .pdf) and returns its plain text content
so it can be fed into the section-generation prompts.
"""
import pathlib


def read_rfp(path: str) -> str:
    p = pathlib.Path(path)
    if not p.exists():
        raise FileNotFoundError(f"RFP file not found: {path}")

    suffix = p.suffix.lower()

    if suffix == ".txt":
        return p.read_text(encoding="utf-8")

    if suffix == ".docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise RuntimeError("python-docx is required to read .docx files: pip install python-docx") from e
        d = docx.Document(str(p))
        parts = [para.text for para in d.paragraphs if para.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        try:
            import pypdf
        except ImportError as e:
            raise RuntimeError("pypdf is required to read .pdf files: pip install pypdf") from e
        reader = pypdf.PdfReader(str(p))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    raise ValueError(f"Unsupported RFP file type: {suffix}. Use .txt, .docx, or .pdf.")
