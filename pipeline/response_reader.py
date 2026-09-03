"""
Phase 2: reads a filled-in Clarification Questions document (client sent it back with the
Answer column completed) and matches each answer back to its requirement_id using the sidecar
mapping file written by clarification_doc_builder.py.
"""
from __future__ import annotations

import json
import pathlib
import re

import docx

_LEADING_INT = re.compile(r"\d+")


def normalize_q_num(raw: str) -> str | None:
    """Extracts the leading integer from a Q# cell's text so minor formatting differences
    between what we generated and what came back don't break matching — '1', 'Q1', '1.', '#1',
    and '1)' all normalize to '1'. Returns None if no digit is found at all."""
    m = _LEADING_INT.search(raw or "")
    return m.group(0) if m else None


def read_answer_rows(filled_docx_path: str) -> list[tuple[str, str, str]]:
    """Returns [(q_num_raw, question_text, answer_text), ...] for every row in the first table
    of the document with a non-blank Answer cell, skipping the header row. `q_num_raw` is
    whatever raw text was in the first cell — callers matching it against a Q# -> requirement_id
    mapping should normalize it first via normalize_q_num()."""
    d = docx.Document(filled_docx_path)
    if not d.tables:
        raise ValueError("No table found in the filled document — expected the Q&A table.")
    table = d.tables[0]

    rows: list[tuple[str, str, str]] = []
    for row in table.rows[1:]:  # skip header
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        q_num, question, answer = cells[0], cells[1], cells[2]
        if not answer:
            continue
        rows.append((q_num, question, answer))
    return rows


def read_client_responses(filled_docx_path: str, mapping_path: str) -> dict[str, str]:
    """
    Returns {requirement_id: client_answer_text}. Skips rows where the Answer column is still
    blank (client didn't answer that one) rather than guessing. Matches each row's Q# against
    the mapping both as-written and normalized (leading digits only), so small formatting
    differences ("1" vs "Q1" vs "1.") don't silently drop an otherwise-good answer.
    """
    mapping = json.loads(pathlib.Path(mapping_path).read_text(encoding="utf-8"))

    responses: dict[str, str] = {}
    for q_num_raw, _question, answer in read_answer_rows(filled_docx_path):
        req_id = mapping.get(q_num_raw) or mapping.get(normalize_q_num(q_num_raw) or "")
        if req_id:
            responses[req_id] = answer

    return responses
