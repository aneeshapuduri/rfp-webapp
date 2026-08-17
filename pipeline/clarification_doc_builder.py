"""
Phase 2: Clarification Questions document generator.

Produces a clean, client-facing .docx with a fillable table (Question # | Question | Answer).
The client types answers directly into the Answer column and sends the file back. A sidecar
JSON file maps each Question # to its internal requirement_id, so when the filled document is
re-uploaded, response_reader.py can match answers back to the exact requirement that raised
them without exposing internal IDs/reasoning to the client.
"""
from __future__ import annotations

import datetime
import json
import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from schema import Phase1Result

NAVY = RGBColor(0x1F, 0x2D, 0x50)
GRAY = RGBColor(0x59, 0x59, 0x59)


def build_clarification_doc(result: Phase1Result, company_name: str, output_path: str) -> dict:
    """
    Returns the Q# -> requirement_id mapping (also written as a sidecar .mapping.json file
    next to output_path) so the caller can persist it against the project record.
    """
    ambiguous = result.get_ambiguous()
    if not ambiguous:
        raise ValueError("No ambiguous requirements — a clarification doc shouldn't be generated.")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"Clarification Questions — {result.project_title}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    run = sub.add_run(f"Prepared by {company_name} for {result.agency}")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = GRAY

    date_p = doc.add_paragraph()
    run = date_p.add_run(datetime.date.today().strftime("%B %d, %Y"))
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "In reviewing this RFP, we identified a small number of items where additional detail "
        "would help us propose the most accurate and cost-effective solution. Please provide "
        "responses in the Answer column below and return this document; we will proceed with "
        "our full proposal upon receipt."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.6), Inches(3.4), Inches(2.5)]
    headers = ["#", "Question", "Answer"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].width = widths[i]
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    mapping: dict[str, str] = {}
    for i, item in enumerate(ambiguous, start=1):
        q_num = str(i)
        mapping[q_num] = item.id
        row = table.add_row().cells
        row[0].text = q_num
        row[1].text = item.clarification_question or ""
        row[2].text = ""  # blank for client to fill in
        for j, cell in enumerate(row):
            cell.width = widths[j]

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        "Thank you — we appreciate the opportunity to clarify these items before finalizing our "
        "proposal."
    ).italic = True

    doc.save(output_path)

    mapping_path = pathlib.Path(output_path).with_suffix(".mapping.json")
    mapping_path.write_text(json.dumps(mapping, indent=2), encoding="utf-8")

    return mapping
